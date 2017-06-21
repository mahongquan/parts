# -*- coding: utf-8 -*-
# 序号
# 选  购  件（推荐用户随机采购，删除客户不采购的部件）
# 牌号
# 标准配置（免费提供）
from docx import Document
from io import BytesIO,StringIO
import logging
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
def treat_zhuji(table):
	m=len(table.rows)
	n=len(table.columns)
	for i in range(m):
		for j in range(n):
			print(table.cell(i,j).text)
def treat_xuangou(table):
	m=len(table.rows)
	n=len(table.columns)
	for i in range(m):
		for j in range(n):
			print(table.cell(i,j).text)
	pass
def treat_biaoyang(table):
	m=len(table.rows)
	n=len(table.columns)
	for i in range(m):
		for j in range(n):
			print(table.cell(i,j).text)	
	pass
def treat_biaopei(table):
	m=len(table.rows)
	n=len(table.columns)
	for i in range(m):
		for j in range(n):
			print(table.cell(i,j).text)	
	pass
def main():	
	document = Document("ON17487-ON-3000技术协议.docx")
	print(document.tables)
	for tbl in document.tables:
		if "序号" in tbl.cell(0,0).text:
			treat_zhuji(tbl)
		elif "选  购  件" in tbl.cell(0,0).text:
			treat_xuangou(tbl)
		elif "牌号" in tbl.cell(0,0).text:
			treat_biaoyang(tbl)
		elif "标准配置" in tbl.cell(0,0).text:
			treat_biaopei(tbl)
if __name__=="__main__":
    main()

