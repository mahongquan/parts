# -*- coding: utf-8 -*-
import sys
import platform
import logging
import json 
from mysite.settings import MEDIA_ROOT,MEDIA_URL
from genDoc.excel_write import *  #证书
from genDoc.docx_write import genPack #装箱单
import genDoc.genLabel          #标签
from genDoc.recordXml import genRecord #调试记录
import re
import django
from django.shortcuts import render_to_response
import time
import os
from django.http import HttpResponse,HttpResponseRedirect
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.hashers import  check_password, make_password
from django.contrib.auth.models import User,Group
from django.core.exceptions import ObjectDoesNotExist#,DoesNotExist
from django.forms.models  import modelform_factory
from django.forms import ModelForm
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.template.context_processors import csrf
import mysite.settings
import datetime
from mysite.parts.models import *
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.models import Group
from django.db.models import Q
from myutil import MyEncoder
import traceback
import xlrd
from django.db import connection,transaction
from docx import Document
# @transaction.commit_on_success
def mylistdir(p,f):
    a=os.listdir(p)
    fs=myfind(a,f)
    return(fs)
def myfind(l,p):
    lr=[];
    #print p
    p1=p.replace(".",r"\.")
    p2=p1.replace("*",".*")
    p2=p2+"$"
    p2="^"+p2
    for a in l:
        #print a
        if  re.search(p2,a,re.IGNORECASE)==None :
           pass
           #print "pass"
        else:
           lr.append(a)
       #print "append"
    return lr
def getIniFile(contact):
    filepath = mysite.settings.MEDIA_ROOT 
    path=os.path.join(filepath, "仪器资料/%s" % (contact.yiqibh))
    try:
        fs=mylistdir(path,"*.ini")
        out="./仪器资料/%s" % (contact.yiqibh)
        if len(fs)>0:
            return  out+"/"+fs[0]
        else:
            pass
    except FileNotFoundError as e:
        logging.info(e)
        pass
    xhp=contact.yiqixinghao.split("-")[0]
    path=os.path.join(filepath,"仪器资料/%s/%s" % (contact.yiqibh,xhp))
    try:
        fs=mylistdir(path,"*.ini")
        out="./仪器资料/%s/%s" % (contact.yiqibh,xhp)
        if len(fs)>0:
            return  out+"/"+fs[0]
        else:
            return None
    except FileNotFoundError as e:
        logging.info(e)
        return None
       
def inItems(item,items):
    inIt=False
    equal=False
    v=None
    for i in range(len(items)):
        if items[i][0]==item[0]:
            inIt=True
            if items[i][2]==item[2]:
                equal=True
            v=items[i]
            items.remove(items[i])
            break
    return(inIt,equal,v)
def printList(items):
    r=[]
    for item in items:
        r1=[]
        for one in item:
            r1.append(str(one))
        r.append(",".join(r1))
    return "\n".join(r)
def bjitems(items,items_chuku):
    #(left,middle,right)bjitems(items,items_chuku)
    logging.info(items)
    left=[]
    equal=[]
    notequal=[]
    for item in items:
        (inIt,equalv,v)=inItems(item,items_chuku)
        if inIt:
            if equalv:
                equal.append(item)
            else:
                notequal.append(item)
                notequal.append(v)
        else:
            left.append(item)
    # print("left")
    # print(printList(left))
    # print("equal")
    # print(printList(equal))
    # print("!equal")
    # print(printList(notequal))
    # print("right")
    # print(printList(items_chuku))
    return(left,notequal,items_chuku)
# def writer(request):
#     # logging.info(request)
#     # output={}
#     # return HttpResponse(json.dumps(output, ensure_ascii=False))
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     r=render_to_response("rest/writer.html",c)
#     return(r)
# @login_required
# def app_users_view(request):
#     start=int(request.GET.get("start","0"))
#     limit=int(request.GET.get("limit","20"))
#     total=User.objects.count()
#     objs =User.objects.all()[start:start+limit]
#     data=[]
#     for rec in objs:
#         data.append({"id":rec.id,"name":str(rec.username),"email":str(rec.email),"first":str(rec.first_name),"last":rec.last_name})
#     logging.info(data)
#     out={"total":total,"data":data}
#     return HttpResponse(json.dumps(out, ensure_ascii=False,cls=MyEncoder))
# @login_required    
# def app_users_update(request):
#     logging.info(request)
#     request2=Request(request,(JSONParser(),))
#     data = request2.DATA['data']
#     id1=data["id"]
#     rec=User.objects.get(id=id1)
#     if data.get("name")!=None:
#         rec.username=data["name"]
#     if data.get("email")!=None:
#         rec.email=data["email"]
#     if data.get("first")!=None:
#         rec.first_name=data["first"]
#     if data.get("last")!=None:
#         rec.last_name=data["last"]
#     rec.save()
#     output={"success":True,"message":"UPDATE new User" +str(rec.id)}
#     output["data"]={"id":rec.id,"name":str(rec.username),"email":str(rec.email),"first":str(rec.first_name),"last":rec.last_name}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
# @login_required    
# def app_users_destroy(request):
#     logging.info(request)
#     request2=Request(request,(JSONParser(),))
#     data = request2.DATA['data']
#     id1=data["id"]
#     rec=User.objects.get(id=id1)
#     rec.delete()
#     output={"success":True,"message":"delete User" +str(rec.id)}
#     output["data"]={"id":id1}#,"name":str(rec.username),"email":str(rec.email),"first":str(rec.first_name),"last":rec.last_name}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
# @login_required
# def app_users_create(request):
#     logging.info(request)
#     request2=Request(request,(JSONParser(),))
#     data = request2.DATA['data']
#     rec=User()#.objects.get(id=id1)
#     if data.get("name")!=None:
#         rec.username=data["name"]
#     if data.get("email")!=None:
#         rec.email=data["email"]
#     if data.get("first")!=None:
#         rec.first_name=data["first"]
#     if data.get("last")!=None:
#         rec.last_name=data["last"]
#     rec.save()
#     output={"success":True,"message":"create new User" +str(rec.id)}
#     output["data"]={"id":rec.id,"name":str(rec.username),"email":str(rec.email),"first":str(rec.first_name),"last":rec.last_name}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
# def index(request):
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     return render_to_response("rest/index.html",c)
# def backbone(request):
#     r=csrf(request)["csrf_token"]
#     logging.info(dir(r))
#     logging.info(r)
#     c={"user":request.user,"csrf_token":r}
#     #c.update(csrf(request))
#     logging.info(dir(c))
#     logging.info(c)
#     r=render_to_response("rest/backbone.html",c)
#     return(r)    
# def restful(request):
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     r=render_to_response("rest/restful.html",c)
#     return(r)
# def jqm(request):
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     r=render_to_response("rest/jqm.html",c)
#     return(r)
# def index_2(request):
#     # logging.info(request)
#     # output={}
#     # return HttpResponse(json.dumps(output, ensure_ascii=False))
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     r=render_to_response("rest/index_2.html",c)
#     return(r) 
# def extjs6(request):
#     c={"user":request.user,"csrf_token":csrf(request)["csrf_token"]}
#     r=render_to_response("rest/extjs6.html",c)
#     return(r)   
@login_required
def item(request):
    logging.info("===================")
    logging.info(request)
    logging.info(dir(request))
    logging.info("------------------")
    #request2=Request(request,(JSONParser(),))
    #logging.info(request2)
    if request.method == 'GET':
        return view_item(request)
    if request.method == 'POST':
        return create_item(request)
    if request.method == 'PUT':
        return update_item(request)
    if request.method == 'DELETE':
        return destroy_item(request)    
# @login_required
# def application(request):
#     logging.info("===================")
#     logging.info(request)
#     logging.info("------------------")
#     request2=request
#     logging.info(request2)
#     if request.method == 'GET':
#         return view(request2)
#     if request.method == 'POST':
#         return create(request2)
#     if request.method == 'PUT':
#         return update(request2)
#     if request.method == 'DELETE':
#         return destroy(request2)
# def view(request):
#     objs=User.objects.all()
#     data=[]
#     for obj in objs:
#         data.append({"id":obj.id,"email":obj.email,"username":obj.username})
#     output={"data":data}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
# def create(request):
#     f=request.META["wsgi.input"]
#     logging.info(dir(f))
#     data="---"
#     logging.info(data)
#     rec=User()
#     rec.username=request.POST["username"]
#     rec.email=request.POST["email"]
#     rec.save()
#     output={"success":True,"message":"Created new User" +str(rec.id)}
#     output["data"]={"id":rec.id,"email":rec.email,"username":rec.username}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
# def update(request):
#     id1=int(request.POST["id"])
#     rec=User.objects.get(id=id1)
#     if request.POST.get("username")!=None:
#         rec.username=request.POST["username"]
#     if request.POST.get("email")!=None:
#         rec.email=request.POST["email"]
#     rec.save()
#     output={"success":True,"message":"Created new User" +str(rec.id)}
#     output["data"]={"id":rec.id,"email":rec.email,"username":rec.username}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))

# def destroy(request):
#     id=request.path.split("/")[-1]
#     id1=int(id)
#     rec=User.objects.get(id=id1)
#     rec.delete()
#     output={"success":True,"message":"OK"}
#     return HttpResponse(json.dumps(output, ensure_ascii=False))
def view_item(request):
    logging.info("here")
    #pack_id=int(request.GET.get("pack"))
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","2000"))
    search=request.GET.get("query",'')
    if search!='':
        total=Item.objects.filter(Q(name__icontains=search)).count()# | Q(bh__icontains=search)
        objs = Item.objects.filter(Q(name__icontains=search)).order_by('name','-id')[start:start+limit]
    else:
        total=Item.objects.count()
        objs = Item.objects.all().order_by('name','-id')[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei,"image":rec.image})
    logging.info(data)
    out={"total":total,"data":data,"success":True}
    return HttpResponse(json.dumps(out, ensure_ascii=False,cls=MyEncoder))
def create_item(request):
    data = json.loads(request.body.decode("utf-8"))
    #logging.info(data)
    #data=request.POST
    logging.info(data)
    requestPOST=data
    rec=Item()
    if requestPOST.get("bh")!=None:
        rec.bh=requestPOST["bh"]
    if requestPOST.get("name")!=None:
        rec.name=requestPOST["name"]
    if requestPOST.get("guige")!=None:
        rec.guige=requestPOST["guige"]
    if requestPOST.get("danwei")!=None:
        rec.danwei=requestPOST["danwei"]
    rec.save()
    output={"success":True,"message":"Created new User" +str(rec.id)}
    output["data"]={"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
def update_item(request):
    requestPOST = json.loads(request.body.decode("utf-8"))
    id1=int(requestPOST["id"])
    rec=Item.objects.get(id=id1)
    if requestPOST.get("bh")!=None:
        rec.bh=requestPOST["bh"]
    if requestPOST.get("name")!=None:
        rec.name=requestPOST["name"]
    if requestPOST.get("guige")!=None:
        rec.guige=requestPOST["guige"]
    if requestPOST.get("danwei")!=None:
        rec.danwei=requestPOST["danwei"]
    rec.save()
    output={"success":True,"message":"update item " +str(rec.id)}
    output["data"]={"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
def destroy_item(request):
    requestPOST = json.loads(request.body.decode("utf-8"))
    id1=int(requestPOST["id"])
    rec=Item.objects.get(id=id1)
    rec.delete()
    output={"success":True,"message":"OK"}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def contact(request):
    logging.info("=contact==========")
    logging.info(request)
    logging.info("------------------")
    if request.method == 'GET':
        return view_contact(request)
    if request.method == 'POST':
        return create_contact(request)
    if request.method == 'PUT':
        return update_contact(request)
    if request.method == 'DELETE':
        return destroy_contact(request)
def updateMethod(request): 
    id1=request.GET.get("id")
    id1=int(id1)
    c=Contact.objects.get(id=id1)       
    c.method=getIniFile(c)
    c.save()
    output={"success":True,"message":"","data":c.json()}
    return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def view_contact(request):
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","20"))
    search=request.GET.get("search",'')
    baoxiang=request.GET.get("baoxiang",'')
    logging.info("search="+search)
    logging.info("baoxiang="+baoxiang)
    if search!='':
        if baoxiang!="": 
            tmp=Contact.objects.filter((Q(hetongbh__icontains=search) | Q(yiqibh__icontains=search) | Q(yonghu__icontains=search)) & Q(baoxiang=baoxiang))
            total=tmp.count()
            objs =tmp.order_by('-yujifahuo_date')[start:start+limit]
        else:
            tmp=Contact.objects.filter(Q(hetongbh__icontains=search)| Q(yonghu__icontains=search) | Q(yiqibh__icontains=search))
            total=tmp.count()
            objs = tmp.order_by('-yujifahuo_date')[start:start+limit]
    else:
        if baoxiang!="":
            tmp=Contact.objects.filter(Q(baoxiang=baoxiang))
            total=tmp.count()
            objs =tmp.order_by('-yujifahuo_date')[start:start+limit]
        else:
            total=Contact.objects.count()
            objs = Contact.objects.order_by('-yujifahuo_date')[start:start+limit]
    data=[]
    for rec in objs:
        data.append(rec.json())
    #logging.info(data)
    output={"total":total,"data":data,"user":request.user.username}
    return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def create_contact(request):
    try:
        logging.info(request.body)
        data = json.loads(request.body.decode("utf-8"))#extjs read data from body
        rec=Contact()
        if data.get("hetongbh")!=None:
            rec.hetongbh=data["hetongbh"]
        if data.get("yujifahuo_date")!=None:
            dt=datetime.datetime.strptime(data["yujifahuo_date"],'%Y-%m-%d')
            rec.yujifahuo_date=dt.date()
        if data.get("yonghu")!=None:
            rec.yonghu=data.get("yonghu")
        if data.get("baoxiang")!=None:
            rec.baoxiang=data.get("baoxiang")
        if data.get("yiqixinghao")!=None:
            rec.yiqixinghao=data.get("yiqixinghao")
        if data.get("yiqibh")!=None:
            rec.yiqibh=data.get("yiqibh")
        if data.get("shenhe")!=None:
            rec.shenhe=data.get("shenhe")
        if data.get("addr")!=None:
            rec.addr=data.get("addr")
        if data.get("channels")!=None:
            rec.channels=data.get("channels")
        if data.get("tiaoshi_date")!=None:
            #rec.tiaoshi_date=datetime.datetime.fromtimestamp(int(data["tiaoshi_date"]))
            dt=datetime.datetime.strptime(data["tiaoshi_date"],'%Y-%m-%d')
            rec.tiaoshi_date=dt.date()
        rec.save()
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]=rec.json()#{"id":rec.id,"shenhe":rec.shenhe,"hetongbh":rec.hetongbh,"yiqibh":rec.yiqibh,"yiqixinghao":rec.yiqixinghao,"yujifahuo_date":rec.yujifahuo_date,"yonghu":rec.yonghu,"baoxiang":rec.baoxiang,"addr":rec.addr,"channels":rec.channels,"tiaoshi_date":rec.tiaoshi_date}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except ValueError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except django.db.utils.IntegrityError as e:
        print(e,dir(e))
        if "hetongbh" in e.args[0]:
            message="合同编号必须唯一，不能重复！"
        elif "yiqibh" in e.args[0]:
            message="仪器编号必须唯一，不能重复！"
        else:
            info = sys.exc_info()
            message=""
            for file, lineno, function, text in traceback.extract_tb(info[2]):
                message+= "%s line:, %s in %s: %s\n" % (file,lineno,function,text)
            message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def update_contact(request):
    try:
        data = json.loads(request.body.decode("utf-8"))#extjs read data from body
        logging.info(data)
        id1=data.get("id")
        id1=int(id1)
        rec=Contact.objects.get(id=id1)
        rec.myupdate(data)
        rec.save()
        output={"success":True,"message":"update Contact " +str(rec.id)}
        output["data"]=rec.json()
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except ValueError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except django.db.utils.IntegrityError as e:
        print(e,dir(e))
        if "hetongbh" in e.args[0]:
            message="合同编号必须唯一，不能重复！"
        elif "yiqibh" in e.args[0]:
            message="仪器编号必须唯一，不能重复！"
        else:
            info = sys.exc_info()
            message=""
            for file, lineno, function, text in traceback.extract_tb(info[2]):
                message+= "%s line:, %s in %s: %s\n" % (file,lineno,function,text)
            message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def destroy_contact(request):
    data = json.loads(request.body.decode("utf-8"))
    id=data.get("id")
    if id!=None:
        try:
            id1=int(id)
            rec=Contact.objects.get(id=id1)
            rec.delete()
            output={"success":True,"message":"OK"}
            return HttpResponse(json.dumps(output, ensure_ascii=False))
        except ObjectDoesNotExist as e:
            output={"success":False,"message":str(e)}
            return HttpResponse(json.dumps(output, ensure_ascii=False))
    else:
        output={"success":False,"message":"OK"}
        return HttpResponse(json.dumps(output, ensure_ascii=False))
def mylogout(request):
    logging.info("logout/////////////////////////////////////////////////")
    logging.info(request)
    logging.info(dir(request.user))
    logout(request)
    output={"success":True,"message":"User"}
    r=HttpResponse(json.dumps(output, ensure_ascii=False))
    logging.info(r)
    return r
def login_index(request):
    output={"success":True,"user":str(request.user),"csrf_token":str(csrf(request)["csrf_token"])}
    r=HttpResponse(json.dumps(output, ensure_ascii=False))
    return(r)
def mylogin(request):
    logging.info("login/////////////////////////////////////////////////")
    logging.info(request.POST)
    request2=request#Request(request,(JSONParser(),))
    data = request2.POST
    username = data['username']
    password = data['password']
    user = authenticate(username=username, password=password)
    if user is None:
        output={"success":False,"message":"No This User"}
    else:
        login(request, user)
        rec=user
        output={"success":True,"message":"User" +str(rec.id)}
        output["data"]={"id":rec.id,"name":str(rec.username),"email":str(rec.email),"first":str(rec.first_name),"last":rec.last_name}
    r=HttpResponse(json.dumps(output, ensure_ascii=False))
    logging.info(r)
    return r
def functions(request):
    logging.info("===================")
    logging.info(request)
    logging.info("------------------")
    request2=Request(request,(JSONParser(),))
    logging.info(request2)
    if request.method == 'GET':
        return view(request2)
    if request.method == 'POST':
        return create(request2)
    if request.method == 'PUT':
        return update(request2)
    if request.method == 'DELETE':
        return destroy(request2)
@login_required
def view_item2(request):
    logging.info("here")
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","20"))
    search=request.GET.get("search",'')
    search_bh=request.GET.get("search_bh",'')
    logging.info("search="+search)
    if search!='':
        if search_bh!='':
            total=Item.objects.filter(name__contains=search).filter(bh__contains=search_bh).count()
            objs = Item.objects.filter(name__contains=search).filter(bh__contains=search_bh)[start:start+limit]
        else:
            total=Item.objects.filter(name__contains=search).count()
            objs = Item.objects.filter(name__contains=search)[start:start+limit]
    else:
        if search_bh!='':
            total=Item.objects.filter(bh__contains=search_bh).count()
            objs = Item.objects.filter(bh__contains=search_bh)[start:start+limit]
        else:
            total=Item.objects.count()
            objs = Item.objects.all()[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei})
    logging.info(data)
    out={"total":total,"data":data}
    return HttpResponse(json.dumps(out, ensure_ascii=False,cls=MyEncoder))
@login_required
def create_item2(request):
    #request=Request(request,(JSONParser(),))
    logging.info(request.POST)
    datas=request.POST["data"]
    logging.info(datas)
    output={"success":True,"message":"Created new User"}
    if type(datas) is list:
        output["data"]=[]
        for data in datas:
            rec=Item()
            rec.name=data["name"]
            rec.bh=data["bh"]
            rec.guige=data["guige"]
            rec.danwei=data["danwei"]
            rec.save()
            output["data"].append({"clientId":data["id"],"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei})
    else:
        data=datas
        rec=Item()
        rec.name=data["name"]
        rec.bh=data["bh"]
        rec.guige=data["guige"]
        rec.danwei=data["danwei"]
        rec.save()
        output["data"]={"clientId":data["id"],"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
    #not batch
    # request=Request(request,(JSONParser(),))
    # logging.info(request.POST)
    # data=request.POST["data"]
 #     rec=Item()
 #     rec.name=data["name"]
 #     rec.bh=data["bh"]
 #     rec.guige=data["guige"]
 #     rec.danwei=data["danwei"]
    # rec.save()
    # output={"success":True,"message":"Created new User" +str(rec.id)}
    # output["data"]={"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei}
    # return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def update_item2(request):
    request=Request(request,(JSONParser(),))
    datas=request.POST["data"]
    output={"success":True,"message":"update item "}
    if type(datas) is list:
        output["data"]=[]
        for data in datas:
            id1=int(data["id"])
            rec=Item.objects.get(id=id1)
            if data.get("bh")!=None:
                 rec.bh=data["bh"]
            if data.get("name")!=None:
                 rec.name=data["name"]
            if data.get("guige")!=None:
                 rec.guige=data["guige"]
            if data.get("danwei")!=None:
                 rec.danwei=data["danwei"]
            rec.save()
            output["data"].append({"clientId":data["id"],"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei})
    else:
        data=datas
        id1=int(data["id"])
        rec=Item.objects.get(id=id1)
        if data.get("bh")!=None:
             rec.bh=data["bh"]
        if data.get("name")!=None:
             rec.name=data["name"]
        if data.get("guige")!=None:
             rec.guige=data["guige"]
        if data.get("danwei")!=None:
             rec.danwei=data["danwei"]
        rec.save()
        output["data"]={"clientId":data["id"],"id":rec.id,"bh":rec.bh,"name":rec.name,"guige":rec.guige,"danwei":rec.danwei}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
    # objs=User.objects.all()
    # data=[]
    # for rec in objs:
    #     data.append({"id":rec.id,"hetongbh":rec.hetongbh,"yujifahuo_date":rec.yujifahuo_date,"yonghu":rec.yonghu,"baoxiang":rec.baoxiang})
    # output={"data":data}
    # return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def destroy_item2(request):
    request=Request(request,(JSONParser(),))
    datas=request.POST["data"]
    if type(datas) is list:
        for data in datas:
            id1=int(data["id"])
            rec=Item.objects.get(id=id1)
            rec.delete()
    else:
        data=datas
        id1=int(data["id"])
        rec=Item.objects.get(id=id1)
        rec.delete()
    output={"success":True,"message":"OK"}
    return HttpResponse(json.dumps(output, ensure_ascii=False))

def organize(request):
    c=RequestContext(request,{})
    c.update(csrf(request))
    r=render_to_response("rest/organizer.html",c)
    return(r)
def geticons(request):
    logging.info("here")
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","20"))
    search=request.GET.get("search",'')
    search_bh=request.GET.get("search_bh",'')
    logging.info("search="+search)
    if search!='':
        if search_bh!='':
            total=Item.objects.filter(name__contains=search).filter(bh__contains=search_bh).count()
            objs = Item.objects.filter(name__contains=search).filter(bh__contains=search_bh)[start:start+limit]
        else:
            total=Item.objects.filter(name__contains=search).count()
            objs = Item.objects.filter(name__contains=search)[start:start+limit]
    else:
        if search_bh!='':
            total=Item.objects.filter(bh__contains=search_bh).count()
            objs = Item.objects.filter(bh__contains=search_bh)[start:start+limit]
        else:
            total=Item.objects.count()
            objs = Item.objects.all()[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"thumb":rec.image.name,"name":rec.name})
    logging.info(data)
    return HttpResponse(json.dumps(data, ensure_ascii=False,cls=MyEncoder))
    # output=[
    #     {
    #         "name": "说明书",
    #         "thumb": "20140326_015.jpg",
    #         "url": "kitchensink",
    #         "type": "Application"
    #     },
    #     {
    #         "name": "坩埚",
    #         "thumb": "0102g001004.JPG",
    #         "url": "twitter",
    #         "type": "Application"
    #     },
    #     {
    #         "name": "真空硅脂",
    #         "thumb": "0102g001009.JPG",
    #         "url": "kiva",
    #         "type": "Application"
    #     },
    #     {
    #         "name": "铜管",
    #         "thumb": "0103a005011.JPG",
    #         "url": "geocongress",
    #         "type": "Application"
    #     }
    # ]
    # return HttpResponse(json.dumps(output, ensure_ascii=False))    
@login_required
def view_pack(request):
    logging.info("here")
    logging.info(request.GET)
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","200"))
    search=request.GET.get("search",'')
    logging.info("search="+search)
    if search!='':
        total=Pack.objects.filter(name__contains=search).count()
        objs = Pack.objects.filter(name__contains=search).order_by('-id')[start:start+limit]
    else:
        total=Pack.objects.count()
        objs = Pack.objects.all().order_by('-id')[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"name":rec.name})
    logging.info(data)
    out={"total":total,"data":data}
    return HttpResponse(json.dumps(out, ensure_ascii=False))
@login_required
def create_pack(request):
    #request=Request(request,(JSONParser(),))
    datas = json.loads(request.body.decode("utf-8"))#extjs read data from body
    #logging.info(request.POST)
    #datas=json.loads(request.POST["data"])
    logging.info(datas)
    datas=datas["data"]
    output={"success":True,"message":"Created new User"}
    if type(datas) is list:
        output["data"]=[]
        for data in datas:
            rec=Pack()
            rec.name=data["name"]
            rec.save()
            output["data"].append({"id":rec.id,"name":rec.name})
    else:
        data=datas
        rec=Pack()
        rec.name=data["name"]
        rec.save()
        output["data"]={"id":rec.id,"name":rec.name}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def update_pack(request):
    request=Request(request,(JSONParser(),))
    datas=request.POST["data"]
    output={"success":True,"message":"update Pack "}
    if type(datas) is list:
        output["data"]=[]
        for data in datas:
            id1=int(data["id"])
            rec=Pack.objects.get(id=id1)
            if data.get("name")!=None:
                rec.name=data["name"]
            rec.save()
            output["data"].append({"clientId":data["id"],"id":rec.id,"name":rec.name})
    else:
        data=datas
        id1=int(data["id"])
        rec=Pack.objects.get(id=id1)
        if data.get("name")!=None:
            rec.name=data["name"]
        rec.save()
        output["data"]={"clientId":data["id"],"id":rec.id,"name":rec.name}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def destroy_pack(request):
    request=Request(request,(JSONParser(),))
    datas=request.POST["data"]
    if type(datas) is list:
        for data in datas:
            id1=int(data["id"])
            rec=Pack.objects.get(id=id1)
            rec.delete()
    else:
        data=datas
        id1=int(data["id"])
        rec=Pack.objects.get(id=id1)
        rec.delete()
    output={"success":True,"message":"OK"}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
#contact pack##################
@login_required
def usepack(request):
    logging.info("===================")
    logging.info(request)
    logging.info("------------------")
    #request2=Request(request,(JSONParser(),))
    #logging.info(request2)
    if request.method == 'GET':
        return view_usepack(request)
    if request.method == 'POST':
        return create_usepack(request)
    if request.method == 'PUT':
        return update_usepack(request)
    if request.method == 'DELETE':
        return destroy_usepack(request)
def view_usepack(request):
    logging.info("view_usepack")
    contact=int(request.GET.get("contact","0"))
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","20"))
    total=UsePack.objects.filter(contact=contact).count()
    objs = UsePack.objects.filter(contact=contact)[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"contact":str(rec.contact.id),"pack":str(rec.pack.id),"hetongbh":rec.contact.hetongbh,"name":rec.pack.name})
    logging.info(data)
    out={"total":total,"data":data}
    return HttpResponse(json.dumps(out, ensure_ascii=False,cls=MyEncoder))
def create_usepack(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    rec=UsePack()
    if(data.get("contact")!=None and data.get("pack")!=None):
        rec.contact=Contact.objects.get(id=int(data["contact"]))
        rec.pack=Pack.objects.get(id=int(data["pack"]))
        logging.info(rec)
        rec.save()
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]={"id":rec.id,"contact":str(rec.contact.id),"pack":str(rec.pack.id),"hetongbh":rec.contact.hetongbh,"name":rec.pack.name}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    else:
        output={"success":False,"message":"No enough parameters"}
        output["data"]={}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def update_usepack(request):
    id1=int(request.POST["id"])
    rec=UsePack.objects.get(id=id1)
    if request.POST.get("contact")!=None:
         rec.contact=request.POST["contact"]
    if request.POST.get("pack")!=None:
         rec.pack=request.POST["pack"]
    rec.save()
    output={"success":True,"message":"update UsePack " +str(rec.id)}
    output["data"]={"id":rec.id,"contact":str(rec.contact.id),"pack":str(rec.pack.id),"hetongbh":rec.contact.hetongbh,"name":rec.pack.name}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
def destroy_usepack(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    id1=int(data["id"])
    rec=UsePack.objects.get(id=id1)
    rec.delete()
    output={"success":True,"message":"OK"}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
#pack##################
def BothPackItem(request):
    if request.method == 'POST':
        return create_BothPackItem(request)
    if request.method == 'PUT':
        return update_BothPackItem(request)
def UsePackEx(request):
    if request.method == 'POST':
        return create_UsePackEx(request)
    if request.method == 'PUT':
        return update_UsePackEx(request)      
def create_UsePackEx(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    logging.info(data)
    if(data.get("name")!=None):
        rec1=Pack()
        rec1.name=data["name"]
        rec1.save()
        
        rec=UsePack()
        rec.pack=rec1
        contactid=int(data.get("contact"))
        contact=Contact.objects.get(id=contactid)
        rec.contact=contact
        rec.save()
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]={"id":rec.id,"name":rec1.name,"contact":rec.contact.id,"pack":rec.pack.id,"hetongbh":rec.contact.hetongbh}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    else:
        output={"success":False,"message":"No enough parameters"}
        output["data"]={}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))        
def update_UsePackEx(request):          
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    id1=int(data["id"])
    rec=UsePack.objects.get(id=id1)
    rec1=rec.item;
    if data.get("name")!=None:
        rec1.name=data["name"]
    rec1.save()
    output={"success":True,"message":"update UsePack " +str(rec.id)}
    output["data"]={"id":rec.id,"name":rec1.name,"contact":rec.contact.id}
    return HttpResponse(json.dumps(output, ensure_ascii=False))          
def create_BothPackItem(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    logging.info(data)
    if(data.get("name")!=None):
        rec1=Item()
        rec1.name=data["name"]
        rec1.save()
        
        rec=PackItem()
        rec.item=rec1
        packid=int(data.get("pack"))
        pack=Pack.objects.get(id=packid)
        rec.pack=pack
        rec.ct=1
        rec.save()
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]={"id":rec.id,"name":rec1.name,"danwei":rec1.danwei,"guige":rec1.guige,"ct":rec1.ct,"bh":rec1.bh,"pack":rec.pack.id}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    else:
        output={"success":False,"message":"No enough parameters"}
        output["data"]={}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))        
def update_BothPackItem(request):          
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    id1=int(data["id"])
    rec=PackItem.objects.get(id=id1)
    rec1=rec.item;
    if data.get("name")!=None:
        rec1.name=data["name"]
    if data.get("guige")!=None:
        rec1.guige=data["guige"]
    if data.get("danwei")!=None:
        rec1.danwei=data["danwei"]
    if data.get("bh")!=None:
        rec1.bh=data["bh"]
    rec1.save()
    recChange=False
    if data.get("quehuo")!=None:
        rec.quehuo=data["quehuo"]
        recChange=True
    if data.get("ct")!=None:
        rec.ct=data["ct"]
        recChange=True
    if recChange:
        rec.save()
    output={"success":True,"message":"update UsePack " +str(rec.id)}
    output["data"]={"quehuo":rec.quehuo,"id":rec.id,"name":rec1.name,"danwei":rec1.danwei,"guige":rec1.guige,"ct":rec.ct,"bh":rec1.bh,"pack":rec.pack.id}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def pack(request):
    logging.info("===================")
    logging.info(request)
    #logging.info("------------------")
    #request2=Request(request,(JSONParser(),))
    #logging.info(request2)
    if request.method == 'GET':
        return view_pack1(request)
    if request.method == 'POST':
        return create_pack1(request)
    if request.method == 'PUT':
        return update_pack1(request)
    if request.method == 'DELETE':
        return destroy_pack1(request)
def view_pack1(request):
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","20"))
    search_bh=request.GET.get("search",'')
    if search_bh!='':
        total=Pack.objects.filter(name__contains=search_bh).count()
        objs =Pack.objects.filter(name__contains=search_bh).order_by("-id")[start:start+limit]
    else:
        total=Pack.objects.count()
        objs =Pack.objects.all().order_by("-id")[start:start+limit]
    #total=Pack.objects.count()
    #objs = Pack.objects.all()[start:start+limit]
    data=[]
    for rec in objs:
        data.append({"id":rec.id,"name":rec.name})
    logging.info(data)
    out={"total":total,"data":data}
    return HttpResponse(json.dumps(out, ensure_ascii=False,cls=MyEncoder))
def create_pack1(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    if(data.get("name")!=None):
        rec=Pack()
        rec.name=data["name"]
        rec.save()
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]={"id":rec.id,"name":rec.name}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    else:
        output={"success":False,"message":"No enough parameters"}
        output["data"]={}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))

def update_pack1(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    id1=int(data["id"])
    rec=Pack.objects.get(id=id1)
    if data.get("name")!=None:
        rec.name=data["name"]
    rec.save()
    output={"success":True,"message":"update UsePack " +str(rec.id)}
    output["data"]={"id":rec.id,"name":rec.name}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
def destroy_pack1(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    #id=request.path.split("/")[-1]
    id1=int(data["id"])
    rec=Pack.objects.get(id=id1)
    rec.delete()
    output={"success":True,"message":"OK"}
    return HttpResponse(json.dumps(output, ensure_ascii=False))
@login_required
def packItem(request):
    if request.method == 'GET':
        return view_packItem(request)
    if request.method == 'POST':
        return create_packItem(request)
    if request.method == 'PUT':
        return update_packItem(request)
    if request.method == 'DELETE':
        return destroy_packItem(request)
def view_packItem(request):
    logging.info("view_packitem")
    contact=int(request.GET.get("pack","0"))
    start=int(request.GET.get("start","0"))
    limit=int(request.GET.get("limit","2000"))
    # search_bh=request.GET.get("search",'')
    # if search_bh!='':
    #     total=PackItem.objects.filter(name__contains=search_bh).count()
    #     objs =PackItem.objects.filter(name__contains=search_bh)[start:start+limit]
    # else:
    #     total=PackItem.objects.count()
    #     objs =PackItem.objects.all()[start:start+limit]
    total=PackItem.objects.filter(pack=contact).count()
    objs =PackItem.objects.filter(pack=contact)[start:start+limit]
    data=[]
    for rec in objs:
        data.append(rec.json())
    output={"data":data,"total":total}
    return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def create_packItem(request):
     #request2=Request(request)
     #logging.info(request.POST)
     #logging.info("body=("+str(request.body)+")")
     data = json.loads(request.body.decode("utf-8"))#extjs read data from body
     logging.info("data=("+str(data)+")")
     rec=PackItem()
     if data.get("pack")!=None:
         rec.pack=Pack.objects.get(id=int(data["pack"]))
     if data.get("itemid")!=None:
         rec.item=Item.objects.get(id=int(data["itemid"]))
     if data.get("ct")!=None:
         rec.ct=float(data.get("ct"))
     rec.save()
     output={"success":True,"message":"Created new User" +str(rec.id)}
     output["data"]=rec.json()
     return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def update_packItem(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    logging.info(data)
    id1=data.get("id")
    if id1!=None:
         id1=int(id1)
         item=Item.objects.get(id=int(data["itemid"]))
         item.bh=data.get("bh")
         item.danwei=data.get("danwei")
         item.name=data.get("name")
         item.guige=data.get("guige")
         item.save()
         rec=PackItem.objects.get(id=id1)
         if data.get("pack")!=None:
             rec.pack=Pack.objects.get(id=int(data["pack"]))
         if data.get("itemid")!=None:
             rec.item=item
         if data.get("ct")!=None:
             rec.ct=float(data.get("ct"))
         if data.get("quehuo")!=None:
             rec.quehuo=data.get("quehuo")
         rec.save()
         output={"success":True,"message":"update Contact " +str(rec.id)}
         output["data"]=rec.json()
         return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    else:
        output={"success":False,"message":"need  id"}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
def destroy_packItem(request):
    data = json.loads(request.body.decode("utf-8"))#extjs read data from body
    id=data.get("id")
    if id!=None:
        try:
            id1=int(id)
            rec=PackItem.objects.get(id=id1)
            rec.delete()
            output={"success":True,"message":"OK"}
            return HttpResponse(json.dumps(output, ensure_ascii=False))
        except ObjectDoesNotExist as e:
            output={"success":False,"message":str(e)}
            return HttpResponse(json.dumps(output, ensure_ascii=False))
    else:
        output={"success":False,"message":"OK"}
        return HttpResponse(json.dumps(output, ensure_ascii=False))
def upload(request):
    # where to store files. Probably best defined in settings.py
    filepath = mysite.settings.MEDIA_ROOT 

    # right, so 'file' is the name of the file upload field
    #print request.FILES
    logging.info(request.FILES)
    f= request.FILES[ 'file' ]
    logging.info(dir(f))
    filename = f.name
    filetype = f.content_type

    #the uploaded data from the file
    #f.open()
    
    data=f.read()

    # the full file path and name
    fullfilepath = os.path.join( filepath, filename )
    # clean up filenames & paths:
    fullfilepath = os.path.normpath( fullfilepath )
    fullfilepath = os.path.normcase( fullfilepath )
    num=1
    newfilename=filename

    while(os.path.exists(fullfilepath)):
        logging.info(num)
        logging.info(fullfilepath)
        newfilename=str(num)+"_"+filename
        fullfilepath=os.path.join( filepath, str(num)+"_"+filename )
        fullfilepath = os.path.normpath( fullfilepath )
        fullfilepath = os.path.normcase( fullfilepath )
        num +=1
    # try to write file to the dir.
    try:
        f = open( fullfilepath, 'wb' ) # Writing in binary mode for windows..?
        f.write( data )
        f.close( )
        res={"success":True, "files":"./"+newfilename}
    except e:
        res={"success":False, "files":str(e)}
    return HttpResponse(json.dumps(res, ensure_ascii=False))
def readChuKUfile(content):
    book = xlrd.open_workbook(file_contents=content)
    table=book.sheets()[0]
    nrows = table.nrows
    ncols = table.ncols
    begin=False
    dan=[]
    for i in range(nrows-9-3):
        #print(i,table.row_values(i)[0])
        cells=table.row_values(9+i)
        dan.append((cells[0],cells[1],cells[4]))#bh,name,ct
    yiqibh=str(int(table.row_values(7)[3]))
    return (dan,yiqibh)
def readBeiliaofile(fn):
    book = xlrd.open_workbook(file_contents=fn)
    table=book.sheets()[0]
    nrows = table.nrows
    ncols = table.ncols
    begin=False
    dan=[]
    for i in range(nrows-7):
        #print(i,table.row_values(i)[0])
        cells=table.row_values(7+i)
        dan.append((cells[0],cells[1],cells[7]))#bh,name,ct
    return dan
def check(request):
    contactid=int(request.POST.get("id"))
    contact=Contact.objects.get(id=contactid)
    #    (fileName,fileType)= QFileDialog.getOpenFileName(None,"Open Excel file", ".","Excel Files ( *.xlsx *.xls)")
    # where to store files. Probably best defined in settings.py
    filepath = mysite.settings.MEDIA_ROOT 

    # right, so 'file' is the name of the file upload field
    #print request.FILES
    logging.info(request.FILES)
    f= request.FILES[ 'file' ]
    logging.info(dir(f))
    filename = f.name
    filetype = f.content_type

    #the uploaded data from the file
    #f.open()
    #data=f.read()
    (items,items2)=contact.huizong()
    r=[]
    for item in items:
        r.append((item.bh,item.name,item.ct))
    for item in items2:
        r.append((item.bh,item.name,item.ct))
    items_chuku=readBeiliaofile(f.read())
    #logging.info(yqbh)
    # if yqbh!=contact.yiqibh:
    #     res={"success":False, "result":""}
    # else:
    (left,notequal,right)=bjitems(r,items_chuku)
        # try to write file to the dir.
    res={"success":True, "result":(left,notequal,right)}
    return HttpResponse(json.dumps(res, ensure_ascii=False))    
def readStandardFile(fn,filename):
    book = xlrd.open_workbook(file_contents=fn)
    table=book.sheets()[0]
    nrows = table.nrows
    ncols = table.ncols
    begin=False
    dan=[]
    for i in range(nrows ):
        cells=table.row_values(i)
        if cells[0]=="其他入库单":
            if not begin:
                begin=True
                onedan=[]
            else:
                #finish
                dan.append(onedan)
                onedan=[]
        else:
            if begin:
                onedan.append(cells)
            else:
                pass
    logging.info(onedan)
    if len(onedan)>0:
        dan.append(onedan)
    rs=[]
    for one in dan:
        r=treatOne(one,filename)
        if r!=None: 
            rs.append(r)
    return rs
def treatOne(rows,fn):
    logging.info(rows)
    r=None
    beizhu=rows[1][7]
    if beizhu[:2]=="CS" or beizhu[:2]=="ON":
        name=rows[1][7]+"_"+fn
        d=Pack.objects.filter(name=name)
        logging.info(d)
        if len(d)>0:
            pass
        else:
            d=Pack()
            d.name=rows[1][7]+"_"+fn
            d.save()
            n=len(rows)
            items=rows[4:n]
            for i in items:
                #i=DanjuItem()
                if i[0]=="合计":
                    break
                print(i[1],i[2],i[3],i[4],i[5])
                items=Item.objects.filter(bh=i[1]).all()
                if len(items)>1:
                    item=items[0]
                else:
                    item=Item()
                item.bh=i[1]
                item.name=str(i[2])+" "+str(i[1])
                item.guige=i[3]
                item.danwei=i[4]
                item.save()
                di=PackItem()
                di.pack=d
                di.item=item
                di.ct=i[5]
                di.save()
            r={"id":d.id,"name":d.name}
    return r
def standard(request):
    # right, so 'file' is the name of the file upload field
    #print request.FILES
    logging.info(request.FILES)
    f= request.FILES[ 'file' ]
    logging.info(dir(f))
    filename = f.name
    filetype = f.content_type
    packs=readStandardFile(f.read(),filename)
    res={"success":True, "result":packs}
    return HttpResponse(json.dumps(res, ensure_ascii=False))   
def readHtFile(fn,filename):
    try:
        document = Document(fn)
        yqxh=document.paragraphs[8].text
        logging.info(yqxh)
        yqxh=yqxh.split("：")[1].strip()
        dt=document.paragraphs[10].text
        dt=dt.split("：")[1].strip()
        if "年" in dt:
            [y,other]=dt.split("年")
            [m,other]=other.split("月")
            [d,other]=other.split("日")
            dt=datetime.date(int(y),int(m),int(d))
        elif "-" in dt:
            [y,m,d]=dt.split("-")
            dt=datetime.date(int(y),int(m),int(d))
        mj=document.paragraphs[11].text
        mj=mj.split("：")[1].strip()
        zz=document.paragraphs[12].text
        zz=zz.split("：")[1].strip()
        if(zz==""):
            zz=mj
        print(yqxh,y,m,d,mj,zz)
        #table0 主机清单
        tbl=document.tables[0]
        data=[]
        for row in tbl.rows:
            one=[]
            for c in row.cells:
                one.append(c.text)
            #print(one)
            data.append(one)
        # logging.info(data[1][1])//CS-2800型 碳硫分析仪主机（最多配置3个通道）

        # yiqi=data[1][1]#yiqixinghao ,channels
        # [xh,other]=yiqi.split(" ")
        # [zj,other]=other.split("（")
        channels=""
        contact=Contact()
        contact.yonghu=zz
        contact.yiqixinghao=yqxh
        contact.baoxiang=""
        contact.channels=channels
        contact.addr=""
        contact.yiqibh=filename
        contact.yujifahuo_date=dt+datetime.timedelta(30)
        n=datetime.datetime.now()
        contact.tiaoshi_date=datetime.date(n.year,n.month,n.day)
        contact.save()
        if len(data)==7:
            pass
        else:
            print("表格行数有改动")
        #table1 选  购  件
        tbl=document.tables[1]
        pack=Pack()
        pack.name="选购"+contact.yiqibh
        pack.save()
        data=[]
        for row in tbl.rows:
            one=[]
            for c in row.cells:
                one.append(c.text)
            #print(one)
            data.append(one)
        for one in data[2:]:#选  购  件
            print(one)
            #print(one[1],one[0],one[2],one[4],one[5])
            if one[0].strip()!="":
                items=Item.objects.filter(bh=one[0]).all()
                if len(items)>1:
                    item=items[0]
                else:
                    item=Item()
            else:
                item=Item()
            item.bh=one[0]
            item.name=one[1]
            item.guige=one[2]
            item.danwei=one[3]
            item.save()
            di=PackItem()
            di.pack=pack
            di.item=item
            di.ct=one[4]
            di.save()
        usepack=UsePack()
        usepack.contact=contact
        usepack.pack=pack;
        usepack.save();
        #table2 标配
        tbl=document.tables[2]
        pack=Pack()
        pack.name="标配"+contact.yiqibh
        pack.save()
        data=[]
        for row in tbl.rows:
            one=[]
            for c in row.cells:
                one.append(c.text)
            #print(one)
            data.append(one)
        for one in data[2:]:
            print(one)
            if one[0].strip()!="":
                items=Item.objects.filter(bh=one[0]).all()
                if len(items)>1:
                    item=items[0]
                else:
                    item=Item()
            else:
                item=Item()
            item.bh=one[0]
            item.name=one[2]
            item.guige=""
            item.danwei=one[3][-1:]
            item.save()
            di=PackItem()
            di.pack=pack
            di.item=item
            di.ct=one[3][:-1]
            di.save()
        usepack=UsePack()
        usepack.contact=contact
        usepack.pack=pack
        usepack.save()
        rec=contact
        output={"success":True,"message":"Created new User" +str(rec.id)}
        output["data"]={"id":rec.id,"shenhe":rec.shenhe,"hetongbh":rec.hetongbh,"yiqibh":rec.yiqibh,"yiqixinghao":rec.yiqixinghao,"yujifahuo_date":rec.yujifahuo_date,"yonghu":rec.yonghu,"baoxiang":rec.baoxiang,"addr":rec.addr,"channels":rec.channels,"tiaoshi_date":rec.tiaoshi_date}
        return  HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except ValueError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except django.db.utils.IntegrityError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s\n" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))        
def ht(request):
    # right, so 'file' is the name of the file upload field
    #print request.FILES
    logging.info(request.FILES)
    f= request.FILES[ 'file' ]
    logging.info(dir(f))
    filename = f.name
    filetype = f.content_type
    return readHtFile(f,filename)
    #return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))  
def year12(request):
    logging.info("chart")
    baoxiang=request.GET.get("baoxiang")
    end_date=datetime.datetime.now()
    start_date=datetime.datetime(end_date.year-12,1,1,0,0,0)
    cursor = connection.cursor()            #获得一个游标(cursor)对象
    #更新操作
    start_date_s=start_date.strftime("%Y-%m-%d")
    end_date_s=end_date.strftime("%Y-%m-%d")
    if baoxiang==None:
        cmd="select strftime('%Y',tiaoshi_date) as year,count(id) from parts_contact  where tiaoshi_date between '"+start_date_s+"' and '"+end_date_s+"' group by year"
    else:
        cmd="select strftime('%Y',tiaoshi_date) as year,count(id) from parts_contact  where baoxiang like '"+baoxiang+"'  and tiaoshi_date between '"+start_date_s+"' and '"+end_date_s+"' group by year"            
    logging.info(cmd)
    cursor.execute(cmd)    #执行sql语句
    raw = cursor.fetchall()                 #返回结果行 或使用 #raw = cursor.fetchall()
    lbls=[]
    values=[]
    for one in raw:
        lbls.append(one[0]+"年")
        values.append(one[1])
    res={"success":True, "lbls":lbls,"values":values}
    return HttpResponse(json.dumps(res, ensure_ascii=False))     
def month12(request):
    logging.info("chart")
    baoxiang=request.GET.get("baoxiang")
    end_date=datetime.datetime.now()
    start_date=datetime.datetime(end_date.year-1,1,1,0,0,0)
    cursor = connection.cursor()            #获得一个游标(cursor)对象
    #更新操作
    start_date_s=start_date.strftime("%Y-%m-%d")
    end_date_s=end_date.strftime("%Y-%m-%d")
    if baoxiang==None:
        cmd="select strftime('%Y-%m',tiaoshi_date) as month,count(id) from parts_contact  where tiaoshi_date between '"+start_date_s+"' and '"+end_date_s+"' group by month"
    else:
        cmd="select strftime('%Y-%m',tiaoshi_date) as month,count(id) from parts_contact  where baoxiang like '"+baoxiang+"'  and tiaoshi_date between '"+start_date_s+"' and '"+end_date_s+"' group by month"            
    logging.info(cmd)
    cursor.execute(cmd)    #执行sql语句
    raw = cursor.fetchall()                 #返回结果行 或使用 #raw = cursor.fetchall()
    lbls=[]
    values=[]
    for one in raw:
        lbls.append(one[0]+"月")
        values.append(one[1])
    res={"success":True, "lbls":lbls,"values":values}
    return HttpResponse(json.dumps(res, ensure_ascii=False))      
def copypack(request):
    try:
        logging.info(request.POST)
        oldid=int(request.POST.get('oldid'))
        newname=request.POST.get('newname')
        logging.info(oldid)
        logging.info(newname)
        old=None
        new=None
        try:
            old=Pack.objects.get(id=oldid) 
        except ObjectDoesNotExist as e:
            pass
        try:
            new=Pack.objects.get(name=newname) 
        except ObjectDoesNotExist as e:
            new=Pack()
            new.name=newname
            new.save()
        #copy items
        content=""
        if old==None:
            content="old is None"
        else:
            for pi in old.packitem_set.all():
                n=PackItem()
                n.pack=new
                n.item=pi.item
                n.ct=pi.ct
                n.save()
            content="复制成功！"
        res={"success":True, "message":content}
        return HttpResponse(json.dumps(res, ensure_ascii=False))   
    except ValueError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))
    except django.db.utils.IntegrityError as e:
        info = sys.exc_info()
        message=""
        for file, lineno, function, text in traceback.extract_tb(info[2]):
            message+= "%s line:, %s in %s: %s\n" % (file,lineno,function,text)
        message+= "** %s: %s" % info[:2]
        output={"success":False,"message":message}
        return HttpResponse(json.dumps(output, ensure_ascii=False,cls=MyEncoder))            
def showcontact(request):
    #print request.GET
    contact_id=request.GET["id"]
    c=Contact.objects.get(id=contact_id)
    dic={}
    dic["contact"]=c
    (items,items2)=c.huizong()
    dic["items"]=items
    if len(items2)==0:
        items2=None
    dic["items2"]=items2
    totalct=0
    for i in items:
        totalct +=i.ct
    dic["totalct"]=totalct
    dic["totalid"]=len(items)
    return HttpResponse(json.dumps(dic, ensure_ascii=False,cls=MyEncoder))     
def allfile(request):
    #try:
        contact_id=request.GET["id"]
        c=Contact.objects.get(id=contact_id)
        
        outfilename=c.yiqixinghao+"_"+c.yonghu
        outfilename=outfilename[0:30]
        dir1="证书_"+outfilename
        #
        #p="d:/parts/media/仪器资料/"+c.yiqibh
        p=os.path.join(MEDIA_ROOT,"仪器资料/"+c.yiqibh)
        #证书
        # dir1=p+"/"+outfilename
        # logging.info(dir1)
        if not os.path.exists(p):
            os.makedirs(p)
        # file1=dir1+"/证书数据表"+EXTNAME
        # if not os.path.exists(file1):
        #     fullfilepath = os.path.join(MEDIA_ROOT,"t_证书数据表"+EXTNAME)
        #     data=genShujubiao(c,fullfilepath)
        #     open(file1,"wb").write(data)
        file2=p+"/"+c.hetongbh+"_"+c.yonghu+"_证书"+EXTNAME
        if not os.path.exists(file2):
            data2=getJiaoZhunFile(c)
            open(file2,"wb").write(data2)
        file3=p+"/"+outfilename+"_装箱单.docx"
        if not os.path.exists(file3):
            if c.yujifahuo_date<datetime.datetime.now().date():
                c.yujifahuo_date=datetime.datetime.now().date()
                c.save()
            fullfilepath = os.path.join(MEDIA_ROOT,"t_装箱单.docx")
            data_zxd=genPack(c,fullfilepath)
            open(file3,"wb").write(data_zxd)
        file4=p+"/"+"标签.lbx"
        if not os.path.exists(file4):
            data_lbl=genDoc.genLabel.genLabel(c.yiqixinghao,c.yiqibh,c.channels)
            open(file4,"wb").write(data_lbl)
        logging.info(c.method)
        logging.info(type(c.method))
        if c.method!=None:
            try:
                logging.info("here")
                fullfilepath = os.path.join(MEDIA_ROOT,c.method.path)
                logging.info(fullfilepath)
                (data_record,data_xishu)=genRecord(fullfilepath,c)
                file5=p+"/"+c.yiqibh+"调试记录.docx"
                if not os.path.exists(file5):
                    open(file5,"wb").write(data_record)
                file6=p+"/"+"系数.lbx"
                if not os.path.exists(file6):
                    open(file6,"wb").write(data_xishu)
            except ValueError as e:
                logging.info(e)
                try:
                    (data_record,data_xishu)=genRecord("",c)
                    file5=p+"/"+c.yiqibh+"调试记录.docx"
                    if not os.path.exists(file5):
                        open(file5,"wb").write(data_record)
                except ValueError as e:
                    logging.info(e)
                    pass
            except:
                traceback.print_exc()
                logging.info("except")
        logging.info(p)
        pf=platform.system()
        if pf=="Linux":
            os.system("xdg-open "+p)
        elif  pf.split("-")[0]=="CYGWIN_NT":
            os.system('cygstart "'+p+'"')
        else:
            os.system('start '+p)
        out={"success":True}
        return HttpResponse(json.dumps(out, ensure_ascii=False))
    # except:
    #     message=""
    #     info = sys.exc_info()
    #     for file, lineno, function, text in traceback.extract_tb(info[2]):
    #         message+= "%s line:, %s in %s: %s\n" % (file,lineno,function,text)
    #     message+= "** %s: %s" % info[:2]
    #     out={"success":False,"message":message}
    #     return HttpResponse(json.dumps(out, ensure_ascii=False))
def folder(request):
    contact_id=request.GET["id"]
    c=Contact.objects.get(id=contact_id)
    #p="d:/parts/media/仪器资料/"+c.yiqibh
    p=os.path.join(MEDIA_ROOT,"仪器资料/"+c.yiqibh)
    logging.info(p)
    
    if not os.path.exists(p):
        os.makedirs(p)
    pf=platform.system()
    logging.info(pf)
    if pf=="Linux":
        os.system("xdg-open "+p)
    elif  pf.split("-")[0]=="CYGWIN_NT":
        os.system('cygstart "'+p+'"')
    else:
        os.system('start '+p)
    out={"success":True}
    return HttpResponse(json.dumps(out, ensure_ascii=False))    
