# -*- coding: utf-8 -*-

import os
import sys
import codecs
import django
from genDoc.recordXml import genRecord
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "mysite.settings")
django.setup()
from mysite.parts.models import *
from docx import Document
from io import BytesIO,StringIO
import logging
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from lxml import etree as ET
document=None
def changeGrid2(tbl,rowv,colv,value):
    tbl.cell(rowv,colv).text=value
def setCell(column1,value,teshu):
    column1.text=value
    if teshu:
        # column1.paragraphs[0].style.font.underline=True
        column1.paragraphs[0].style=document.styles['me_underline']
    # logging.info(column1.style)
    # logging.info(dir(column1))
def getGrid(tbl,rowv,colv):
    tbl_childs=tbl.getchildren()
    row1=tbl_childs[rowv+2]
    columns=row1.getchildren()[1:]
    print(columns)
    column1=columns[colv]
    cell=column1.getchildren()[1]
    paras=cell.getchildren()[1:]
    r=[]
    for p in paras:
        r.append(p.getchildren()[1].text)
    return "".join(r)
def getElement(chanels,first):
    eles=first.split("(")
    ele=eles[0]#2C
    if ele[0]=="2":
        chanels.append("L"+ele[1])
        chanels.append("H"+ele[1])
    else:
        ele_set=eles[1][:-1]#移去括号
        print(ele_set)
        if ele_set[0]=="高":
            chanels.append("H"+ele[1])
        else:
            chanels.append("L"+ele[1])
    pass
def myformat_date(d):
    return "%d-%d-%d" %    (d.year,d.month,d.day)
def addItem(items,item):
    find=False
    for i in items:
        if i.id==item.id:
            i.ct +=item.ct
            find=True
            break
    if not find:
        items.append(item)
    return items
def border(cell):
    tc=cell._tc
    ns='xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:ve="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
    e=ET.fromstring("""
        <w:tcBorders %s>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        </w:tcBorders>""" % ns)
    tc.tcPr.insert(1,e)
def borderCells(cells):
    for cell in cells:
        border(cell)
def genPack(contact,fn):
    global document
    document = Document(fn)
    tbl=document.tables[0]
    changeGrid2(tbl,0,1,contact.yonghu)
    changeGrid2(tbl,1,1,contact.yiqixinghao)
    changeGrid2(tbl,2,1,contact.yiqibh)
    changeGrid2(tbl,3,1,contact.baoxiang)
    changeGrid2(tbl,4,1,contact.shenhe)
    changeGrid2(tbl,5,1,myformat_date(contact.yujifahuo_date))
    changeGrid2(tbl,6,1,contact.hetongbh)
    tbl=document.tables[1]#主机清单
    if contact.channels==None:
        changeGrid2(tbl,1,2,contact.yiqixinghao)
    else:
        changeGrid2(tbl,1,2,contact.yiqixinghao+" "+contact.channels)
    if contact.yiqixinghao[0]=='C':
        changeGrid2(tbl,2,1,"")
        pass
        #changeGrid2(tbl,2,3,"1根")
    elif contact.yiqixinghao[0]=='O':
        changeGrid2(tbl,2,3,"3根")
    elif contact.yiqixinghao[0]=='N':
        changeGrid2(tbl,2,3,"3根")
    else :
        changeGrid2(tbl,2,1,"")
    tbl=document.tables[2]#配件清单
    print(tbl.cell(0,0).text)
    items=[]
    items2=[]
    style = document.styles.add_style('me_underline', WD_STYLE_TYPE.PARAGRAPH)
    style.font.underline=True
    # for cp in contact.usepack_set.all():
    #     for pi in cp.pack.packitem_set.all():
    #         pi.item.ct=pi.ct
    #         if not pi.quehuo:
    #             items=addItem(items,pi.item)
    #         else:
    #             items2=addItem(items2,pi.item)
    (items,items2)=contact.huizong()
    for item in items:

        tmp_row=tbl.add_row()
        
        # if item.leijia:
        #     tmp_row.style.font.italic = True
        columns= tmp_row.cells
        if item.bh!=None:
            setCell(columns[0],item.bh,item.leijia)
        setCell(columns[1],item.name,item.leijia)
        if item.guige!=None:
            setCell(columns[2],item.guige,item.leijia)
        if item.danwei==None:
            item.danwei=""
        if item.leijia:
            setCell(columns[3],str(item.ct)+item.danwei,item.leijia)
        else:
            setCell(columns[3],str(item.ct)+item.danwei,item.leijia)

        if item.beizhu!=None:
            setCell(columns[5],item.beizhu,item.leijia)
        else:
            setCell(columns[5],"",item.leijia)
        setCell(columns[4],"",item.leijia)
    if len(items2)>0:
        document.add_page_break()
        p=document.add_paragraph('短缺物资清单')
        print(dir(p))
        p.alignment=1
        # # print(dir(p.paragraph_format))
        r=p.runs[0]
        # print(r.font)
        # print(r.font.size)
        # print(dir(r))
        r.font.size=203200 #
        # print(r.style.name)
        # print(r.part)
        # print(dir(r.part))


        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        borderCells(hdr_cells)
        hdr_cells[0].text = '编号'

        hdr_cells[1].text = '名称'
        hdr_cells[2].text = '规格'
        hdr_cells[3].text = '数量'
        for item in items2:
            row_cells = table.add_row().cells
            borderCells(row_cells)
            if item.bh!=None:
                row_cells[0].text = item.bh
            row_cells[1].text = item.name
            if item.guige!=None:
                row_cells[2].text = item.guige
            row_cells[3].text = str(item.ct)+item.danwei
    s=BytesIO()
    document.save(s)
    s.seek(0)
    data=s.read()
    return data
def test():
    document = Document("t_装箱单.docx")
    return document
if __name__ == "__main__":
    contact=Contact.objects.get(id=323)
    data=genPack(contact,"media/t_装箱单.docx")
    f=open("out.docx","wb")
    f.write(data)
    f.close()