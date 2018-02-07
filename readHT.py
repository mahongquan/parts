# -*- coding: utf-8 -*-
import os
import sys
import codecs
import django
import datetime
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "mysite.settings")
django.setup()
from mysite.parts.models import *
from docx import Document
# yonghu = models.CharField(max_length=30,verbose_name="用户单位")#用户单位
#     addr = models.CharField(max_length=30,verbose_name="客户地址",null=True,blank=True)#用户单位
#     channels = models.CharField(max_length=30,verbose_name="通道配置",null=True,blank=True)#用户单位
#     yiqixinghao=models.CharField(max_length=30,verbose_name="仪器型号")#仪器型号
#     yiqibh=models.CharField(unique=True,max_length=30,verbose_name="仪器编号")#仪器编号
#     baoxiang =  models.CharField(max_length=30,verbose_name="包箱")#包箱
#     shenhe =  models.CharField(max_length=30,verbose_name="审核")#审核
#     yujifahuo_date = models.DateField(verbose_name="预计发货时间")#预计发货时间
#     tiaoshi_date = models.DateField(null=True,blank=True,verbose_name="调试时间",default=datetime.datetime.now)#预计发货时间
#     hetongbh=models.CharField(max_length=30,verbose_name="合同编号")#合同编号
#     method=models.FileField(null=True,blank=True,verbose_name="方法")
def read(fn):
    document = Document(fn)
    i=0
    yqxh=document.paragraphs[8].text
    yqxh=yqxh.split("：")[1].strip()
    dt=document.paragraphs[10].text
    dt=dt.split("：")[1].strip()
    [y,other]=dt.split("年")
    [m,other]=other.split("月")
    [d,other]=other.split("日")
    dt=datetime.date(int(y),int(m),int(d))
    mj=document.paragraphs[11].text
    mj=mj.split("：")[1].strip()
    zz=document.paragraphs[12].text
    zz=zz.split("：")[1].strip()
    if(zz==""):
        zz=mj
    print(yqxh,y,m,d,mj,zz)
    #table0 主机清单
    tbl=document.tables[0]
    data=[]
    for row in tbl.rows:
        one=[]
        for c in row.cells:
            one.append(c.text)
        #print(one)
        data.append(one)
    yiqi=data[1][1]#yiqixinghao ,channels
    [xh,other]=yiqi.split(" ")
    [zj,other]=other.split("（")
    channels=other
    contact=Contact()
    contact.yonghu=zz
    contact.yiqixinghao=yqxh
    contact.baoxiang=""
    contact.channels=channels
    contact.addr=""
    contact.yiqibh=str(datetime.datetime.now())
    contact.yujifahuo_date=dt+datetime.timedelta(30)
    contact.save()
    if len(data)==7:
        pass
    else:
        print("表格行数有改动")
    #table1 选  购  件
    tbl=document.tables[1]
    pack=Pack()
    pack.name="选购"+contact.yiqibh
    pack.save()
    data=[]
    for row in tbl.rows:
        one=[]
        for c in row.cells:
            one.append(c.text)
        #print(one)
        data.append(one)
    for one in data[2:]:#选  购  件
        print(one[1],one[0],one[2],one[4],one[5])
        if one[0].strip()!="":
            items=Item.objects.filter(bh=one[0]).all()
            if len(items)>1:
                item=items[0]
            else:
                item=Item()
        else:
            item=Item()
        item.bh=one[0]
        item.name=one[1]
        item.guige=one[2]
        item.danwei=one[3]
        item.save()
        di=PackItem()
        di.pack=pack
        di.item=item
        di.ct=one[4]
        di.save()
    usepack=UsePack()
    usepack.contact=contact
    usepack.pack=pack;
    usepack.save();
    #table2 标配
    tbl=document.tables[2]
    pack=Pack()
    pack.name="标配"+contact.yiqibh
    pack.save()
    data=[]
    for row in tbl.rows:
        one=[]
        for c in row.cells:
            one.append(c.text)
        #print(one)
        data.append(one)
    for one in data[2:]:
        print(one)
        if one[0].strip()!="":
            items=Item.objects.filter(bh=one[0]).all()
            if len(items)>1:
                item=items[0]
            else:
                item=Item()
        else:
            item=Item()
        item.bh=one[0]
        item.name=one[2]
        item.guige=""
        item.danwei=one[3][-1:]
        item.save()
        di=PackItem()
        di.pack=pack
        di.item=item
        di.ct=one[3][:-1]
        di.save()
    usepack=UsePack()
    usepack.contact=contact
    usepack.pack=pack;
    usepack.save();    
if __name__=="__main__":
    #print(datetime.datetime.now())
    read("ON17470-O-3000技术协议.docx")