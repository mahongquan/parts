# -*- coding: utf-8 -*-
import os
import sys
import codecs
import django
from genDoc.recordXml import genRecord
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "mysite.settings")
django.setup()
from mysite.parts.models import *
from docx import Document
from io import BytesIO,StringIO
import logging
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from lxml import etree as ET
document=None
def changeGrid2(tbl,rowv,colv,value):
    tbl.cell(rowv,colv).text=value
def setCell(column1,value,teshu):
    column1.text=value
    if teshu:
        # column1.paragraphs[0].style.font.underline=True
        column1.paragraphs[0].style=document.styles['me_underline']
    # logging.info(column1.style)
    # logging.info(dir(column1))
def getGrid(tbl,rowv,colv):
    tbl_childs=tbl.getchildren()
    row1=tbl_childs[rowv+2]
    columns=row1.getchildren()[1:]
    print(columns)
    column1=columns[colv]
    cell=column1.getchildren()[1]
    paras=cell.getchildren()[1:]
    r=[]
    for p in paras:
        r.append(p.getchildren()[1].text)
    return "".join(r)
def getElement(chanels,first):
    eles=first.split("(")
    ele=eles[0]#2C
    if ele[0]=="2":
        chanels.append("L"+ele[1])
        chanels.append("H"+ele[1])
    else:
        ele_set=eles[1][:-1]#移去括号
        print(ele_set)
        if ele_set[0]=="高":
            chanels.append("H"+ele[1])
        else:
            chanels.append("L"+ele[1])
    pass
def myformat_date(d):
    return "%d-%d-%d" %    (d.year,d.month,d.day)
def addItem(items,item):
    find=False
    for i in items:
        if i.id==item.id:
            i.ct +=item.ct
            find=True
            break
    if not find:
        items.append(item)
    return items
def border(cell):
    tc=cell._tc
    ns='xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:ve="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
    e=ET.fromstring("""
        <w:tcBorders %s>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        </w:tcBorders>""" % ns)
    tc.tcPr.insert(1,e)
def borderCells(cells):
    for cell in cells:
        border(cell)
def genPack(fn):
    global document
    document = Document(fn)
    tbl=document.tables[0]
    # print(dir(tbl))
    # print(len(tbl.rows))
    col=0
    for j in range(4):
        t=tbl.cell(0,j).text
        if  "仪器编号" in t:
            col=j
            break
    print(len(tbl.rows))
    x=range(len(tbl.rows))
    for i in x:
        # print(i)
        # if i==0:
        #     print("continue")
        #     continue
        # else:
        print(tbl.cell(i,col).text)
if __name__ == "__main__":
    # contact=Contact.objects.get(id=323)
    for i in range(12):
        print(i)
        data=genPack(r"D:\2018年工作总结\%d月份工作总结\马红权-2018年%d月工作总结.docx" % (i+1,i+1) )
    # f=open("out.docx","wb")
    # f.write(data)
    # f.close()