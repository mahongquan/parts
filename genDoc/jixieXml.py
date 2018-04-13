# -*- coding: utf-8 -*-
from docx import Document
from genDoc.iniXml import getFromIni
import random
import os
from mysite.settings import MEDIA_ROOT
import re
import logging
from io import BytesIO
#import getkb
import genDoc.genLabel
# import oa
def mylistdir(p,f):
    a=os.listdir(p)
    fs=myfind(a,f)
    return(fs)
def myfind(l,p):
    lr=[];
    #print p
    p1=p.replace(".",r"\.")
    p2=p1.replace("*",".*")
    p2=p2+"$"
    for a in l:
        #print a
        if  re.search(p2,a,re.IGNORECASE)==None :
           pass
           #print "pass"
        else:
           lr.append(a)
       #print "append"
    return lr
def getRound(stdconc):
    a=str(stdconc)
    at=a.find(".")
    return len(a)-1-at
def geteleRsd(stdconc):
    return 0.05
def genOne(stdconc):
    roundws=getRound(stdconc)
    stdconc=float(stdconc)
    sd=pow(10,-roundws)
    #print(stdconc,sd)
    test=round(stdconc-2*sd+random.random()*4*sd,roundws)
    fmt="%0."+str(roundws)+"f"
    #print(fmt,test)
    test_str=fmt % test
    return (test_str)
def changeGrid(tbl,rowv,colv,value):
    print("================changeGrid")
    print(tbl.cell(rowv,colv).text)
    print(value)
    tbl.cell(rowv,colv).text=value
def changeGridMulti(tbl,rowv,colv,value):
    changeGrid(tbl,rowv,colv,"".join(value))
def setCell(column1,value):
    column1.text=value
def getchangdu(chanels):
    if "LO" in chanels:
        lofmt="LO(%d)mm" % 100
    else:
        lofmt="LO()mm"
    if "HO" in chanels:
        hofmt="  HO(%d)mm" % 10
    else:
        hofmt="  HO()mm"
    return lofmt+hofmt
def chanelNums(chanels):
    cs=0
    ss=0
    if "LC" in chanels:
        cs+=1
    if "HC" in chanels:
        cs+=1
    if "LS" in chanels:
        ss+=1
    if "HS" in chanels:
        ss+=1
    return (cs,ss)
def getElement(chanels,first):
    print(first)
    if first=="":
        return
    eles=first.split("(")
    ele=eles[0]#2C
    if ele[0]=="2":
        chanels.append("L"+ele[1])
        chanels.append("H"+ele[1])
    else:
        ele_set=eles[1][:-1]#移去括号
        print(ele_set)
        if ele_set[0]=="高":
            chanels.append("H"+ele[1])
        else:
            chanels.append("L"+ele[1])
    pass    
def getchannels(peizhi):
    print(peizhi)
    elements=peizhi.split("+")
    chanels=[]
    first=elements[0]
    getElement(chanels,first)
    if len(elements)==1:#单元素
        pass
    else:#two elements
        second=elements[1]
        getElement(chanels,second)
    return chanels
def genRecord(fn,c):
    yiqibh=c.yiqibh
    yiqixinghao=c.yiqixinghao
    factors=[]
    if yiqixinghao[0]=="C":
        data=genRecordCS(c)
    else:
        data=genRecordONH(c)
    return data
def genRecordONH(contact):
    #tree = ET.parse(os.path.join(MEDIA_ROOT,'ONH机械.xml'))
    tree = Document(os.path.join(MEDIA_ROOT,'ONH机械.docx'))
    tbls=tree.tables
    ps=tree.paragraphs
    ps[1].runs[-1].text=baoxiang
    ps[4].runs[-1].text=baoxiang
    # i=0
    # for p in tree.paragraphs:
    #     print(i)
    #     print(p.text)
    #     i+=1
    # input("hi")
    tbl=tbls[0]
    changeGrid(tbl,0,1,yiqibh)#
    changeGrid(tbl,0,3,yiqixinghao)#
    tbl=tbls[2]
    #(k,b)=getkb.getkb()
    #changeGrid(tbl,16,1,"电流"+"PID"+"参数：斜率为"+str(k)+"；截距为"+str(b))#
    tbl=tbls[3]
    if "LO" in chanels:
        changeGrid(tbl,1,3,genOne("1.3"))#低碳前置放大电压
    else:
        changeGrid(tbl,1,3,"-")#低碳前置放大电压
    if "HO" in chanels:
        changeGrid(tbl,2,3,genOne("1.3"))#高碳前置放大电压
    else:
        changeGrid(tbl,2,3,"-")#高碳前置放大电压
    #自动调零范围
    if "LO" in chanels:
        changeGrid(tbl,3,3,"".join([genOne("-7.0"),"~",genOne("7.0"),"V"]))#低碳前置放大电压
    else:
        changeGrid(tbl,3,3,"".join(["-","~","-","V"]))#低碳前置放大电压
    if "HO" in chanels:
        changeGrid(tbl,4,3,"".join([genOne("-7.0"),"~",genOne("7.0"),"V"]))#低碳前置放大电压
    else:
        changeGrid(tbl,4,3,"".join(["-","~","-","V"]))#低碳前置放大电压
    if "LN" in chanels:
        changeGrid(tbl,5,3,"".join([genOne("-7.0"),"~",genOne("7.0"),"V"]))#低碳前置放大电压
    else:
        changeGrid(tbl,5,3,"".join(["-","~","-","V"]))#低碳前置放大电压
    if "HN" in chanels:
        changeGrid(tbl,6,3,"".join([genOne("-7.0"),"~",genOne("7.0"),"V"]))#低碳前置放大电压
    else:
        changeGrid(tbl,6,3,"".join(["-","~","-","V"]))#低碳前置放大电压
    changdu=getchangdu(chanels)
    changeGrid(tbl,7,2,changdu)
    # #6、    线性化调试结果
    tbl=tbls[4]
    print("===aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa========")
    print(chanels)
    print(factors)
    if factors!=None:
        if "LO" in chanels:
            changeGrid(tbl,1,2,"%0.1f" % (factors["低氧"][0]))#低碳线性化系数
            changeGrid(tbl,1,3,"%0.3f" % (factors["低氧"][1]))#低碳线性化系数
            changeGrid(tbl,1,4,"%0.3f" % (factors["低氧"][2]))#低碳线性化系数
        else:
            changeGrid(tbl,1,2,"-")#低碳线性化系数
            changeGrid(tbl,1,3,"-")#低碳线性化系数
            changeGrid(tbl,1,4,"-")#低碳线性化系数
        if "HO" in chanels:
            changeGrid(tbl,2,2,"%0.1f" % (factors["高氧"][0]))#高碳线性化系数
            changeGrid(tbl,2,3,"%0.3f" % (factors["高氧"][1]))#高碳线性化系数
            changeGrid(tbl,2,4,"%0.3f" % (factors["高氧"][2]))#高碳线性化系数
        else:
            changeGrid(tbl,2,2,"-")#高碳线性化系数
            changeGrid(tbl,2,3,"-")#高碳线性化系数
            changeGrid(tbl,2,4,"-")#高碳线性化系数
        if "LN" in chanels:
            changeGrid(tbl,3,2,"∞")#低硫线性化系数
            changeGrid(tbl,3,3,"%0.3f" % (factors["低氮"][1]))
            changeGrid(tbl,3,4,"%0.3f" % (factors["低氮"][2]))
        else:
            changeGrid(tbl,3,2,"-")
            changeGrid(tbl,3,3,"-")
            changeGrid(tbl,3,4,"-")
    s=BytesIO()
    tree.save(s)
    s.seek(0)
    data=s.read()
    return data
def genRecordCS(contact):
    #tree = ET.parse(os.path.join(MEDIA_ROOT,'CS机械.xml'))
    tree = Document(os.path.join(MEDIA_ROOT,'CS机械.docx'))
    print(tree.paragraphs)
    c2=tree.paragraphs[1]#仪器编号
    print(c2.runs)
    c2.runs[0].text="仪器编号:"+str(contact.hongwai)
    c2.runs[1].text=""
    c2.runs[2].text=""
    c2.runs[3].text=""

    #红外检测器调试检查
    tbl=tree.tables[0]
    for i in range(len(tbl.rows))[1:]:
        changeGrid(tbl,i,1,str(contact.dianqi))#红外光源电压
        changeGrid(tbl,i,2,str(contact.yujifahuo_date))#红外光源电压
    s=BytesIO()
    tree.save(s)
    s.seek(0)
    data=s.read()
    return data
if __name__=="__main__":
    print(getpath.getpath())
    fs=mylistdir(".","*装箱单.xml")
    if len(fs)>0:
        fn=fs[0].split("_")[0]
        genRecord(fn,fs[0])
        oa.mainUpload(fn)
    else:
        print("未发现文件‘*装箱单.xml’")