# -*- coding: utf-8 -*-
from docx import Document
from io import BytesIO,StringIO
import logging
def changeGrid2(tbl,rowv,colv,value):
    tbl.cell(rowv,colv).text=value
def setCell(column1,value):
    column1.text=value
def getGrid(tbl,rowv,colv):
    tbl_childs=tbl.getchildren()
    row1=tbl_childs[rowv+2]
    columns=row1.getchildren()[1:]
    print(columns)
    column1=columns[colv]
    cell=column1.getchildren()[1]
    paras=cell.getchildren()[1:]
    r=[]
    for p in paras:
        r.append(p.getchildren()[1].text)
    return "".join(r)
def getElement(chanels,first):
    eles=first.split("(")
    ele=eles[0]#2C
    if ele[0]=="2":
        chanels.append("L"+ele[1])
        chanels.append("H"+ele[1])
    else:
        ele_set=eles[1][:-1]#移去括号
        print(ele_set)
        if ele_set[0]=="高":
            chanels.append("H"+ele[1])
        else:
            chanels.append("L"+ele[1])
    pass
def myformat_date(d):
    return "%d-%d-%d" %    (d.year,d.month,d.day)
def addItem(items,item):
    find=False
    for i in items:
        if i.id==item.id:
            i.ct +=item.ct
            find=True
            break
    if not find:
        items.append(item)
    return items
def genPack(contact,fn):
    document = Document(fn)
    tbl=document.tables[0]
    logging.info(dir(tbl))
    changeGrid2(tbl,0,1,contact.yonghu)
    changeGrid2(tbl,1,1,contact.yiqixinghao)
    changeGrid2(tbl,2,1,contact.yiqibh)
    changeGrid2(tbl,3,1,contact.baoxiang)
    changeGrid2(tbl,4,1,contact.shenhe)
    changeGrid2(tbl,5,1,myformat_date(contact.yujifahuo_date))
    changeGrid2(tbl,6,1,contact.hetongbh)
    tbl=document.tables[1]#主机清单
    if contact.channels==None:
        changeGrid2(tbl,1,2,contact.yiqixinghao)
    else:
        changeGrid2(tbl,1,2,contact.yiqixinghao+" "+contact.channels)
    if contact.yiqixinghao[0]=='C':
        changeGrid2(tbl,2,3,"1根")
    else:
        changeGrid2(tbl,2,3,"3根")
    tbl=document.tables[2]#配件清单
    print(tbl.cell(0,0).text)
    items=[]
    items2=[]
    for cp in contact.usepack_set.all():
        for pi in cp.pack.packitem_set.all():
            pi.item.ct=pi.ct
            if not pi.quehuo:
                items=addItem(items,pi.item)
            else:
                items2=addItem(items2,pi.item)
    for item in items:
        columns= tbl.add_row().cells
        setCell(columns[0],item.bh)
        setCell(columns[1],item.name)
        setCell(columns[2],item.guige)
        setCell(columns[3],str(item.ct)+item.danwei)
        if item.beizhu!=None:
            setCell(columns[5],item.beizhu)
        else:
            setCell(columns[5],"")
        setCell(columns[4],"")
    # if len(items2)>0:
    #     document.add_page_break()
    #     table = document.add_table(rows=1, cols=4)
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = '编号'
    #     hdr_cells[1].text = '名称'
    #     hdr_cells[2].text = '规格'
    #     hdr_cells[3].text = '数量'
    #     for item in items2:
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = item.bh
    #         row_cells[1].text = item.name
    #         row_cells[2].text = item.guige
    #         row_cells[3].text = str(item.ct)+item.danwei
    s=BytesIO()
    document.save(s)
    s.seek(0)
    data=s.read()
    return data
def genQue(contact,fn):
    document = Document(fn)
    tbl=document.tables[0]
    logging.info(dir(tbl))
    changeGrid2(tbl,0,1,contact.yonghu)
    changeGrid2(tbl,1,1,contact.yiqixinghao)
    changeGrid2(tbl,2,1,contact.yiqibh)
    changeGrid2(tbl,3,1,contact.baoxiang)
    changeGrid2(tbl,4,1,contact.shenhe)
    changeGrid2(tbl,5,1,myformat_date(contact.yujifahuo_date))
    changeGrid2(tbl,6,1,contact.hetongbh)
    tbl=document.tables[1]
    print(tbl.cell(0,0).text)
    items=[]
    items2=[]
    for cp in contact.usepack_set.all():
        for pi in cp.pack.packitem_set.all():
            pi.item.ct=pi.ct
            if not pi.quehuo:
                items=addItem(items,pi.item)
            else:
                items2=addItem(items2,pi.item)
    for item in items2:
        columns= tbl.add_row().cells
        setCell(columns[0],item.bh)
        setCell(columns[1],item.name)
        setCell(columns[2],item.guige)
        setCell(columns[3],str(item.ct)+item.danwei)
        if item.beizhu!=None:
            setCell(columns[5],item.beizhu)
        else:
            setCell(columns[5],"")
        setCell(columns[4],"")
    s=BytesIO()
    document.save(s)
    s.seek(0)
    data=s.read()
    return data

if __name__=="__main__":
    print(genPack("4111533499"))