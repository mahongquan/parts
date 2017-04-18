# -*- coding: utf-8 -*-
from docx import Document
from io import BytesIO,StringIO
import logging
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.text.font import Font
from lxml import etree as ET
document = Document("gen.docx")
# print(len(document.paragraphs))
# for p in document.paragraphs:
#     if p.text=="配货清单":
#         print(dir(p))
#         # print(p.text,"("+p.style.name+")")
#         # print(p.style.font.size)
#         # print(dir(p._p))
#         # print(p._p.tag)
#         #print(p.alignment,dir(p.alignment),type(p.alignment))
#         r=p.runs[0]
#         #print(dir(r))
#         # print(r.font.size)
#         # r.font.size=r.font.size*2
#         f=r.font
#print(len(document.tables))
tbl=document.tables[1]
print(dir(tbl.cell(0,0)._tc.xml))
print(tbl.cell(0,0)._tc.xml)
# print(tbl.style)
# print(dir(tbl.style))
p=document.add_paragraph('短缺物资清单')
p.alignment=1
r=p.runs[0]
r.font.size=203200#330200 #



table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '编号'
#print(dir(hdr_cells[0]))
#print(dir(hdr_cells[0]._tc))
hdr_cells[1].text = '名称'
hdr_cells[2].text = '规格'
hdr_cells[3].text = '数量'
print(dir(hdr_cells[0]))
def border(cell):
    tc=cell._tc
    ns='xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:ve="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
    e=ET.fromstring("""
        <w:tcBorders %s>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="6" w:space="0" w:color="auto"/>
        </w:tcBorders>""" % ns)
    tc.tcPr.insert(1,e)
#tc.insert(0,e)
# print(table.style)
# print(dir(table.style))
# print(table.style.name)
document.save("gen2.docx")
